from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Create document
doc = Document()

# Title
title = doc.add_heading('BCP Alumni Management System', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Admin Dashboard Modules & Sub-Modules')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].bold = True
subtitle.runs[0].font.size = Pt(14)

doc.add_paragraph()

# Module data with functionalities
modules = [
    {
        "name": "1. Dashboard",
        "description": "Main dashboard view - Overview of system statistics and quick access to all modules",
        "submodules": []
    },
    {
        "name": "2. Alumni & Records",
        "submodules": [
            ("Registered Users", "View and manage registered alumni accounts", 
             "View users list, Search/Filter by batch/course/status, View user details, Status badges"),
            ("Approvals", "Approve pending alumni registrations", 
             "View pending registrations, Approve/Reject applications, Email notifications, Provider verification"),
            ("Records", "Comprehensive alumni records database", 
             "View all alumni records, Search/Filter by year/course, Edit records, Send tracer surveys, Export data"),
            ("Master List", "Upload and manage master list data", 
             "CSV file upload, Batch import records, View uploaded records, Delete records, Download template"),
        ]
    },
    {
        "name": "3. Career & Jobs",
        "submodules": [
            ("Manage Jobs", "Create and manage job postings", 
             "Create/Edit/Delete job posts, Toggle job status (active/inactive), View applicants list, Update application status, Add departments"),
            ("Placement Logs", "Track job placement activities", 
             "View placement records, Filter by industry/status, View placement details, Export to CSV, Statistics dashboard"),
            ("Status Tracker", "Monitor alumni career status", 
             "View employment status, Filter by status/course/batch, Charts and analytics, Real-time statistics"),
        ]
    },
    {
        "name": "4. Events & Reunions",
        "submodules": [
            ("Event Calendar", "Manage events and calendar", 
             "Create/Edit/Archive events, Calendar view, Upload event images, View RSVP list, Date/Time picker, Filter events"),
            ("Batch Reunions", "Organize batch reunion events", 
             "Create/Edit/Archive reunions, Set max attendees, View RSVP attendees, Upload reunion images, Filter by batch year"),
            ("Event Scheduling Integration", "External event management system integration", 
             "Schedule form submission, Integration with external event management system, View pending requests, Approve/Reject event requests, Status tracking, Notification system"),
        ]
    },
    {
        "name": "5. Communication & Updates",
        "submodules": [
            ("News Feed", "Manage news articles and updates", 
             "Create/Edit/Delete articles, Publish/Unpublish toggle, Category management, Image upload, Search/Filter articles"),
        ]
    },
    {
        "name": "6. Engagement",
        "submodules": [
            ("Feedback & Surveys", "Collect and manage alumni feedback", 
             "View feedback list, Reply to feedback, Update status, Create surveys, Add/Remove questions, View survey responses, Toggle survey status"),
            ("Donations", "Manage donation records", 
             "Create/Edit campaigns, View donations, Verify/Reject donations, Archive/Restore campaigns, Upload campaign images, View donation proofs"),
            ("Financial Collections", "Track financial collections", 
             "Real-time donation tracking, Filter by date/status, Charts and analytics, Export to CSV, Masked amounts for privacy"),
        ]
    },
    {
        "name": "7. Advanced Tools",
        "submodules": [
            ("Analytics", "View data analytics and statistics", 
             "Batch year distribution chart, Course distribution chart, Employment status pie chart, Registration trends, Refresh data"),
            ("Report Generator", "Generate custom reports", 
             "Select report type, Filter by date range, Password-protected PDF generation, Charts in reports, Multiple report templates"),
            ("Train AI", "Train AI chatbot responses", 
             "Upload training CSV, Start training workflow, View training logs, n8n integration, Model accuracy display"),
            ("Audit Trail", "View system activity logs", 
             "View all user actions, Search logs, Filter by action/date/user, Export to CSV, Pagination, Action color coding"),
        ]
    },
]

# Add modules
for module in modules:
    # Module heading
    heading = doc.add_heading(module["name"], level=1)
    
    if "description" in module:
        doc.add_paragraph(module["description"])
    
    if module["submodules"]:
        # Create table with 3 columns now
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Sub-Module'
        hdr_cells[1].text = 'Description'
        hdr_cells[2].text = 'Functionalities'
        
        # Make header bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Data rows
        for sub_name, sub_desc, sub_func in module["submodules"]:
            row_cells = table.add_row().cells
            row_cells[0].text = sub_name
            row_cells[1].text = sub_desc
            row_cells[2].text = sub_func
        
        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(1.5)
            row.cells[1].width = Inches(2)
            row.cells[2].width = Inches(3.5)
    
    doc.add_paragraph()

# Summary section
doc.add_heading('Summary', level=1)
summary = doc.add_paragraph()
summary.add_run('Total Main Modules: ').bold = True
summary.add_run('7\n')
summary.add_run('Total Sub-Modules: ').bold = True
summary.add_run('18')

# Save document
output_path = r'c:\Users\arabino\Downloads\BCP_Admin_Modules_List.docx'
doc.save(output_path)
print(f"Document saved to: {output_path}")
